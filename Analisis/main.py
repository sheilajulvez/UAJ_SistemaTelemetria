import pandas as pd
import json
import os
import io
import sys
from docx import Document

# Para escribir en un buffer y luego volcarlo en el documento
buffer = io.StringIO()
sys.stdout = buffer


def cargar_multiples_trazas(carpeta):
    dataframes = []
    session_n = 1 

    for archivo in os.listdir(carpeta):
        if archivo.endswith('.json') or archivo.endswith('.jsonl'):
            ruta = os.path.join(carpeta, archivo)
            with open(ruta, 'r', encoding='utf-8') as file:
                lineas = file.readlines()
                datos = [json.loads(linea) for linea in lineas if linea.strip()]
                df = pd.DataFrame(datos)
                df['archivo'] = archivo
                df['session_n'] = session_n
                dataframes.append(df)
                session_n += 1
    return pd.concat(dataframes, ignore_index=True)

def analizar(df):
    df['event_timestamp'] = pd.to_datetime(df['event_timestamp'], errors='coerce', utc=True)

    print("Conteo de eventos:")
    conteos = df['event_type'].value_counts()
    print(conteos)

    print(f"\n---\n")


    numero_de_sesiones = len(df[df['event_type'] == 'SessionStart'])


    ends = df[df['event_type'] == 'LevelEnd']
    print(f"Total de niveles completados: {len(ends)}")
    if (len(ends) / numero_de_sesiones) >= config["numero_niveles"]:
        print(f"Todos los jugadores han completado los {config["numero_niveles"]} niveles")
    elif (len(ends) / numero_de_sesiones) < config["numero_niveles"]:
        print(f"Algún jugador no ha conseguido superar los {config["numero_niveles"]} niveles.")


    print(f"\n---\n")


    pauses = df[df['event_type'] == 'Pause']
    print(f"Total de pausas entre las {numero_de_sesiones} sesiones: {len(pauses)}")
    print(f"Número de pausas por nivel: {((len(pauses) / numero_de_sesiones) / float(config["numero_niveles"]))}")


    print(f"\n---\n")


    muertesFall = df[df['event_type'] == 'FallDeath']
    muertesSpike = df[df['event_type'] == 'SpikeDeath']
    muertesSlime = df[df['event_type'] == 'SlimeDeath']
    print(f"Total de muertes entre las {numero_de_sesiones} sesiones: {len(muertesSpike) + len(muertesFall) + len(muertesSlime)}")
    print(f"Muertes por caida: {len(muertesFall)}")
    print(f"Muertes por pincho: {len(muertesSpike)}")
    print(f"Muertes por slime: {len(muertesSlime)}")


    print(f"\n---\n")
          
    if numero_de_sesiones > 0:
        duraciones = []

        for session_n, grupo in df.groupby('session_n'):
            inicio = grupo['event_timestamp'].min()
            fin = grupo['event_timestamp'].max()
            duracion = (fin - inicio).total_seconds()
            duraciones.append(duracion)
            print(f"Tiempo jugado en sesión {session_n}: {duracion:.2f} segundos")

        if duraciones:
            promedio = sum(duraciones) / len(duraciones)
            print(f"Tiempo promedio por sesión: {promedio:.2f} segundos")
    else:
        print("No hay datos válidos para calcular duración de sesiones.")


if __name__ == "__main__":
    print(f"\n---\n")
    
    with open("config.json", "r", encoding="utf-8") as f:
        config = json.load(f)
    carpeta = config["carpeta"]
    df = cargar_multiples_trazas(carpeta)
    analizar(df)
    
    sys.stdout = sys.__stdout__
    contenido = buffer.getvalue()

    doc = Document()
    doc.add_heading("Informe automático de análisis – AMONRA", level=1)
    doc.add_paragraph(contenido)
    doc.save("informe_resultados.docx")

    print("\n✅ Se ha guardado un informe en 'informe_resultados.docx'\n")
